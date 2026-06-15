from __future__ import annotations

from abc import ABC, abstractmethod

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class BaseRenderer(ABC):
    """
    Template-agnostic base for all doc-maker renderers.

    Subclasses must implement:
      - render_chapter_opener(meta)
      - render_directive(name, text)

    Everything else (markdown block/inline rendering, XML helpers) is shared
    and inherited by all templates.
    """

    def __init__(self, doc: Document, md_parser):
        self.doc  = doc
        self._md  = md_parser
        self._align = WD_ALIGN_PARAGRAPH.LEFT

    # ── abstract interface ────────────────────────────────────────────────────

    @abstractmethod
    def render_chapter_opener(self, meta: dict) -> None: ...

    @abstractmethod
    def render_directive(self, name: str, text: str) -> None: ...

    # ── shared block rendering ────────────────────────────────────────────────

    def render_tokens(self, tokens: list) -> None:
        for tok in tokens:
            self.render_block(tok)

    def render_block(self, tok: dict) -> None:
        t = tok["type"]

        if t == "heading":
            self.doc.add_heading(
                self._text(tok.get("children", [])), level=tok["level"]
            )

        elif t in ("paragraph", "block_text"):
            p = self.doc.add_paragraph()
            p.alignment = self._align
            self._inline(p, tok.get("children", []))

        elif t in ("newline", "blank_line"):
            pass

        elif t == "thematic_break":
            self._rule()

        elif t == "block_code":
            self._code_block(tok.get("text", ""), tok.get("info", "") or "")

        elif t == "block_quote":
            for child in tok.get("children", []):
                if child["type"] == "paragraph":
                    p = self.doc.add_paragraph(style="Intense Quote")
                    self._inline(p, child.get("children", []))

        elif t == "list":
            self._list(tok.get("children", []), tok.get("ordered", False))

        elif t == "table":
            self._table(tok)

    # ── list ──────────────────────────────────────────────────────────────────

    def _list(self, items: list, ordered: bool) -> None:
        style = "List Number" if ordered else "List Bullet"
        for item in items:
            if item["type"] != "list_item":
                continue
            inline_children: list = []
            sub_lists: list = []
            for child in item.get("children", []):
                if child["type"] == "list":
                    sub_lists.append(child)
                elif child["type"] in ("paragraph", "block_text"):
                    inline_children.extend(child.get("children", []))
                else:
                    inline_children.append(child)
            p = self.doc.add_paragraph(style=style)
            p.alignment = self._align
            self._inline(p, inline_children)
            for sub in sub_lists:
                self._list(sub.get("children", []), sub.get("ordered", False))

    # ── table ─────────────────────────────────────────────────────────────────

    def _table(self, tok: dict) -> None:
        head = next((c for c in tok.get("children", []) if c["type"] == "table_head"), None)
        body = next((c for c in tok.get("children", []) if c["type"] == "table_body"), None)
        head_cells = head.get("children", []) if head else []
        body_rows  = body.get("children", []) if body else []
        all_rows   = ([{"children": head_cells}] if head_cells else []) + body_rows
        if not all_rows:
            return
        cols  = max(len(r.get("children", [])) for r in all_rows)
        table = self.doc.add_table(rows=len(all_rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(all_rows):
            is_head = r_idx == 0 and bool(head_cells)
            for c_idx, cell_tok in enumerate(row.get("children", [])):
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                self._inline(p, cell_tok.get("children", []))
                if is_head:
                    for run in p.runs:
                        run.bold = True

    # ── inline rendering ──────────────────────────────────────────────────────

    def _inline(self, p, tokens: list, bold: bool = False, italic: bool = False) -> None:
        for tok in tokens:
            self._inline_tok(p, tok, bold, italic)

    def _inline_tok(self, p, tok: dict, bold: bool, italic: bool) -> None:
        t = tok["type"]

        if t == "text":
            r = p.add_run(tok.get("text", ""))
            r.bold   = bold
            r.italic = italic

        elif t == "strong":
            self._inline(p, tok.get("children", []), bold=True, italic=italic)

        elif t == "emphasis":
            self._inline(p, tok.get("children", []), bold=bold, italic=True)

        elif t == "codespan":
            r = p.add_run(tok.get("text", ""))
            r.bold           = bold
            r.italic         = italic
            r.font.name      = "Courier New"
            r.font.size      = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

        elif t == "strikethrough":
            for child in tok.get("children", []):
                r = p.add_run(self._text([child]))
                r.font.strike = True

        elif t == "image":
            from ..core.preprocessor import PLACEHOLDER_RE
            src = tok.get("src", "")
            alt = tok.get("alt", "") or ""
            m   = PLACEHOLDER_RE.match(src)
            if m:
                self._image_placeholder(p, m.group(1), alt)

        elif t == "link":
            r = p.add_run(self._text(tok.get("children", [])))
            r.bold           = bold
            r.italic         = italic
            r.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
            r.underline      = True

        elif t == "softlinebreak":
            p.add_run(" ")

        elif t == "linebreak":
            p.add_run("\n")

    # ── XML / style helpers ───────────────────────────────────────────────────

    def _text(self, tokens: list) -> str:
        parts = []
        for tok in tokens:
            if tok["type"] in ("text", "codespan"):
                parts.append(tok.get("text", ""))
            elif "children" in tok:
                parts.append(self._text(tok["children"]))
        return "".join(parts)

    def _rgb(self, hex_str: str) -> RGBColor:
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _shade_p(self, p, fill: str) -> None:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill.upper())
        pPr.append(shd)

    def _borders_p(self, p, *,
                   left: tuple[str, str] | None   = None,
                   right: tuple[str, str] | None  = None,
                   top: tuple[str, str] | None    = None,
                   bottom: tuple[str, str] | None = None) -> None:
        """Add any combination of borders to a paragraph. Each arg: (hex_color, sz_str)."""
        pPr = p._p.get_or_add_pPr()
        for existing in pPr.findall(qn("w:pBdr")):
            pPr.remove(existing)
        bdr = OxmlElement("w:pBdr")
        for side_name, value in [("left", left), ("right", right), ("top", top), ("bottom", bottom)]:
            if value:
                color, sz = value
                el = OxmlElement(f"w:{side_name}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    sz)
                el.set(qn("w:space"), "4")
                el.set(qn("w:color"), color.upper())
                bdr.append(el)
        if len(bdr):
            pPr.append(bdr)

    def _space_p(self, p, before: int = 0, after: int = 0) -> None:
        """Set paragraph spacing (values in dxa — twentieths of a point)."""
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        if before:
            sp.set(qn("w:before"), str(before))
        if after:
            sp.set(qn("w:after"), str(after))
        pPr.append(sp)

    def _left_border(self, p, color: str, sz: str = "18") -> None:
        pPr = p._p.get_or_add_pPr()
        bdr  = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    sz)
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), color.upper())
        bdr.append(left)
        pPr.append(bdr)

    def _indent_p(self, p, left: int = 360) -> None:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        pPr.append(ind)

    def _rule(self, color: str = "AAAAAA", sz: str = "6") -> None:
        p   = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    sz)
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color.upper())
        bdr.append(bot)
        pPr.append(bdr)

    def _page_break(self) -> None:
        p   = self.doc.add_paragraph()
        run = p.add_run()
        br  = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def _code_block(self, code: str, lang: str = "") -> None:
        p   = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F5F5F5")
        pPr.append(shd)
        r = p.add_run(code)
        r.font.name = "Courier New"
        r.font.size = Pt(9)

    def _image_placeholder(self, p, name: str, alt: str = "") -> None:
        label = alt or name
        r = p.add_run(f"[ IMAGE: {label} ]")
        r.font.name      = "Arial"
        r.font.size      = Pt(10)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r.font.italic    = True
