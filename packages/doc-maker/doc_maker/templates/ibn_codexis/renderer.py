from __future__ import annotations

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ...core.base_renderer import BaseRenderer
from ...core.preprocessor import PLACEHOLDER_RE, IMG_INLINE_RE
from .theme import (
    PRIMARY, ACCENT, BODY, CAPTION,
    BORDER_SOFT, BORDER_H2,
    H2_PT, H3_PT, H4_PT, H5_PT,
    BOX_FONT_HDR, BOX_FONT_BODY,
    CODE_BG, CODE_FONT,
    FIGURE_BG, FIGURE_FONT,
    TABLE_HDR,
    BOX_DIRECTIVES, LABEL_DIRECTIVES,
)


class IbnCodexisRenderer(BaseRenderer):
    """
    IBN Codexis Computer Science textbook style.

    Coloured boxes match the source .docx exactly:
      KEY TERM (green), WORKED EXAMPLE (olive), ACTIVITY (teal),
      REMEMBER (navy), EXAM TIP (brown), DID YOU KNOW (purple).

    Headings use borders + colour instead of Word heading styles.
    Data tables have a dark blue header row.
    Code blocks use a VS Code dark theme.
    Document header/footer is injected in render_chapter_opener().
    """

    # ── chapter opener ────────────────────────────────────────────────────────

    def render_chapter_opener(self, meta: dict) -> None:
        chapter_num = meta.get("chapter", "")
        title       = meta.get("title", "")
        subtitle    = meta.get("subtitle", "")
        grade       = meta.get("grade", "")

        self._setup_header_footer(chapter_num, title, grade)

        # "CHAPTER N" small label
        if chapter_num:
            p = self.doc.add_paragraph()
            r = p.add_run(f"CHAPTER  {chapter_num}")
            r.font.size      = Pt(10)
            r.font.color.rgb = self._rgb(ACCENT)
            r.font.all_caps  = True
            r.font.bold      = True
            self._space_p(p, before=200, after=40)

        # Full-width rule
        self._rule(ACCENT, sz="24")

        # Chapter title
        if title:
            p = self.doc.add_paragraph()
            r = p.add_run(title)
            r.bold           = True
            r.font.size      = Pt(22)
            r.font.color.rgb = self._rgb(PRIMARY)
            self._space_p(p, before=80, after=40)

        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph()
            r = p.add_run(subtitle)
            r.italic         = True
            r.font.size      = Pt(11)
            r.font.color.rgb = self._rgb(CAPTION)
            self._space_p(p, before=0, after=160)

        self.doc.add_paragraph()

    def _setup_header_footer(self, chapter_num, title, grade) -> None:
        section = self.doc.sections[0]

        # ── header ────────────────────────────────────────────────────────────
        hdr = section.header
        hdr_p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hdr_p.clear()

        r1 = hdr_p.add_run("IBN CODEXIS INSTITUTE")
        r1.bold           = True
        r1.font.size      = Pt(8)
        r1.font.color.rgb = self._rgb(PRIMARY)

        r2 = hdr_p.add_run("    |    ")
        r2.font.size      = Pt(8)
        r2.font.color.rgb = self._rgb(CAPTION)

        chapter_label = f"Chapter {chapter_num} — {title}" if chapter_num and title else title or ""
        if grade:
            chapter_label += f"  ({grade})"
        r3 = hdr_p.add_run(chapter_label)
        r3.font.size      = Pt(8)
        r3.font.color.rgb = self._rgb(CAPTION)

        self._borders_p(hdr_p, bottom=(ACCENT, "4"))

        # ── footer ────────────────────────────────────────────────────────────
        ftr = section.footer
        ftr_p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        ftr_p.clear()

        grade_label = f"ICT {grade}" if grade else "ICT"
        r4 = ftr_p.add_run(f"IBN Codexis Institute  |  {grade_label}  |  Page ")
        r4.font.size      = Pt(8)
        r4.font.color.rgb = self._rgb(CAPTION)
        self._add_page_number(ftr_p)

    def _add_page_number(self, p) -> None:
        r = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r._r.append(fld_begin)

        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = " PAGE "
        r2._r.append(instr)

        r3 = p.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r3._r.append(fld_sep)

        p.add_run("1")

        r5 = p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r5._r.append(fld_end)

    # ── heading rendering ─────────────────────────────────────────────────────

    def render_block(self, tok: dict) -> None:
        if tok["type"] == "heading":
            self._render_heading(tok)
        else:
            super().render_block(tok)

    def _render_heading(self, tok: dict) -> None:
        level = tok["level"]
        text  = self._text(tok.get("children", []))

        if level == 2:
            p = self.doc.add_paragraph()
            self._borders_p(p, left=(ACCENT, "24"), bottom=(BORDER_H2, "10"))
            self._indent_p(p, 120)
            self._space_p(p, before=300, after=120)
            r = p.add_run(text)
            r.bold           = True
            r.font.size      = Pt(H2_PT)
            r.font.color.rgb = self._rgb(ACCENT)

        elif level == 3:
            p = self.doc.add_paragraph()
            self._borders_p(p, bottom=(BORDER_SOFT, "4"))
            self._space_p(p, before=220, after=80)
            r = p.add_run(text)
            r.bold           = True
            r.font.size      = Pt(H3_PT)
            r.font.color.rgb = self._rgb(ACCENT)

        elif level == 4:
            p = self.doc.add_paragraph()
            self._space_p(p, before=160, after=60)
            r = p.add_run(text)
            r.bold           = True
            r.font.size      = Pt(H4_PT)
            r.font.color.rgb = self._rgb(PRIMARY)

        elif level == 5:
            p = self.doc.add_paragraph()
            self._space_p(p, before=80, after=40)
            r = p.add_run(text)
            r.bold           = True
            r.font.size      = Pt(H5_PT)
            r.font.color.rgb = self._rgb(PRIMARY)

        else:
            p = self.doc.add_paragraph()
            r = p.add_run(text)
            r.bold           = True
            r.font.color.rgb = self._rgb(BODY)

    # ── data table (dark blue header row) ────────────────────────────────────

    def _table(self, tok: dict) -> None:
        super()._table(tok)
        if not self.doc.tables:
            return
        tbl = self.doc.tables[-1]
        if not tbl.rows:
            return
        for cell in tbl.rows[0].cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  TABLE_HDR)
            tcPr.append(shd)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold           = True
                    run.font.color.rgb = self._rgb("FFFFFF")

    # ── code block (dark VS Code style) ──────────────────────────────────────

    def _code_block(self, code: str, lang: str = "") -> None:
        p = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._shade_p(p, CODE_BG)
        self._indent_p(p, 160)
        self._space_p(p, before=80, after=80)
        r = p.add_run(code)
        r.font.name      = "Courier New"
        r.font.size      = Pt(9)
        r.font.color.rgb = self._rgb(CODE_FONT)

    # ── directive dispatcher ──────────────────────────────────────────────────

    def render_directive(self, name: str, text: str, title: str = "") -> None:
        if name == "figure":
            self._render_figure(text)
            return
        if name in ("table-note", "code-note"):
            self._render_caption(text)
            return

        box_cfg = BOX_DIRECTIVES.get(name)
        if box_cfg:
            label, hdr_fill, body_fill = box_cfg
            self._box(label, title, hdr_fill, body_fill, text)
            return

        lbl_cfg = LABEL_DIRECTIVES.get(name)
        if lbl_cfg:
            label, color, pt = lbl_cfg
            self._label_block(label, color, pt, text)
            return

        # Unknown directive — render as plain markdown
        tokens = self._md(text)
        if tokens:
            self.render_tokens(tokens)

    # ── coloured box ─────────────────────────────────────────────────────────

    def _box(self, label: str, title: str, hdr_fill: str, body_fill: str, text: str) -> None:
        """Dark header paragraph + light body paragraphs, matching source .docx box tables."""
        # Header
        header_text = f"  {label}   {title}" if title else f"  {label}"
        p = self.doc.add_paragraph()
        self._shade_p(p, hdr_fill)
        self._space_p(p, before=80, after=0)
        r = p.add_run(header_text)
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = self._rgb(BOX_FONT_HDR)
        r.font.all_caps  = True

        # Body
        for tok in (self._md(text) or []):
            self._block_in_box(tok, body_fill)

        # Closing spacer (gives visual bottom padding)
        p_close = self.doc.add_paragraph()
        self._shade_p(p_close, body_fill)
        self._space_p(p_close, before=0, after=80)

        self.doc.add_paragraph()

    def _block_in_box(self, tok: dict, fill: str) -> None:
        t = tok["type"]

        if t in ("paragraph", "block_text"):
            p = self.doc.add_paragraph()
            self._shade_p(p, fill)
            self._indent_p(p, 200)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            self._inline(p, tok.get("children", []))

        elif t == "list":
            self._box_list(tok, fill)

        elif t == "block_code":
            self._code_block(tok.get("text", ""), tok.get("info", "") or "")

        elif t not in ("newline", "blank_line"):
            self.render_block(tok)

    def _box_list(self, tok: dict, fill: str, depth: int = 0, counter: list | None = None) -> None:
        """Render a list inside a box as plain shaded paragraphs with manual numbering.

        Avoids Word list styles so shading is preserved and counter always starts at 1.
        """
        ordered = tok.get("ordered", False)
        if counter is None:
            counter = [0]

        for item in tok.get("children", []):
            if item["type"] != "list_item":
                continue

            inline: list = []
            sub_lists: list = []
            for child in item.get("children", []):
                if child["type"] == "list":
                    sub_lists.append(child)
                elif child["type"] in ("paragraph", "block_text"):
                    inline.extend(child.get("children", []))
                else:
                    inline.append(child)

            p = self.doc.add_paragraph()
            self._shade_p(p, fill)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            # Manual indent + hanging first-line, matching a real list appearance
            left = 200 + depth * 360
            hang = 280
            pPr  = p._p.get_or_add_pPr()
            ind  = OxmlElement("w:ind")
            ind.set(qn("w:left"),    str(left + hang))
            ind.set(qn("w:hanging"), str(hang))
            pPr.append(ind)

            if ordered:
                counter[0] += 1
                p.add_run(f"{counter[0]}.")
            else:
                p.add_run("•")  # bullet •

            p.add_run("\t")
            self._inline(p, inline)

            for sl in sub_lists:
                self._box_list(sl, fill, depth + 1)

    # ── flat label block ──────────────────────────────────────────────────────

    def _label_block(self, label: str, color: str, pt: float, text: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(label)
        r.bold           = True
        r.font.size      = Pt(pt)
        r.font.color.rgb = self._rgb(color)
        self._space_p(p, before=160, after=80)

        for tok in (self._md(text) or []):
            self.render_block(tok)

        self.doc.add_paragraph()

    # ── figure placeholder ────────────────────────────────────────────────────

    def _render_figure(self, text: str) -> None:
        lines = [ln for ln in text.strip().splitlines() if ln.strip()]
        img_match = None
        caption_parts: list[str] = []

        for line in lines:
            m = IMG_INLINE_RE.search(line.strip())
            if m and img_match is None:
                img_match = m
            else:
                caption_parts.append(line.strip())

        # Grey placeholder box
        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_p(p_img, FIGURE_BG)
        self._space_p(p_img, before=60, after=60)

        if img_match:
            alt, src = img_match.group(1), img_match.group(2)
            ph = PLACEHOLDER_RE.match(src)
            if ph:
                r = p_img.add_run("\n  [ IMAGE / DIAGRAM PLACEHOLDER ]\n")
                r.bold           = True
                r.font.size      = Pt(9)
                r.font.color.rgb = self._rgb(FIGURE_FONT)
                if alt:
                    r2 = p_img.add_run(f"  {alt}\n")
                    r2.italic        = True
                    r2.font.size     = Pt(8)
                    r2.font.color.rgb = self._rgb(FIGURE_FONT)
            else:
                r = p_img.add_run(f"[ IMAGE: {alt or src} ]")
                r.italic         = True
                r.font.color.rgb = self._rgb(FIGURE_FONT)

        # Caption line below the box (bold italic, PRIMARY colour)
        if caption_parts:
            p_cap = self.doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_cap.add_run(" ".join(caption_parts))
            r.bold           = True
            r.italic         = True
            r.font.size      = Pt(9)
            r.font.color.rgb = self._rgb(PRIMARY)
            self._space_p(p_cap, before=40, after=160)

        self.doc.add_paragraph()

    # ── table / code caption ─────────────────────────────────────────────────

    def _render_caption(self, text: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(text.strip())
        r.italic         = True
        r.font.size      = Pt(8.5)
        r.font.color.rgb = self._rgb(CAPTION)
        self._space_p(p, before=40, after=160)
