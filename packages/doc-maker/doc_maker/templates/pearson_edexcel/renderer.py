from __future__ import annotations

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...core.base_renderer import BaseRenderer
from ...core.preprocessor import PLACEHOLDER_RE, IMG_INLINE_RE
from .theme import DIRECTIVE_CONFIG, ACCENT_COLOR, SUBTITLE_COLOR, COVER_FILL


class PearsonEdexcelRenderer(BaseRenderer):
    """
    Renders documents in the Pearson Edexcel student textbook style.
    Supports all 30 elements defined in the Edexcel template.
    """

    # ── chapter opener ────────────────────────────────────────────────────────

    def render_chapter_opener(self, meta: dict) -> None:
        chapter_num = meta.get("chapter", "")
        title       = meta.get("title", "")
        subtitle    = meta.get("subtitle", "")
        cover_image = meta.get("cover_image", "")

        if cover_image:
            m = PLACEHOLDER_RE.match(str(cover_image))
            if m:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._shade_p(p, COVER_FILL)
                r = p.add_run(f"\n  [ IMAGE: {m.group(1)} ]\n")
                r.font.name      = "Arial"
                r.font.size      = Pt(10)
                r.font.color.rgb = self._rgb(ACCENT_COLOR)
                r.font.italic    = True

        if chapter_num:
            p = self.doc.add_paragraph()
            r = p.add_run(f"CHAPTER  {chapter_num}")
            r.font.size      = Pt(10)
            r.font.color.rgb = self._rgb(ACCENT_COLOR)
            r.font.all_caps  = True
            r.font.bold      = True

        self._rule(ACCENT_COLOR, sz="18")

        if title:
            p = self.doc.add_paragraph()
            r = p.add_run(title)
            r.bold           = True
            r.font.size      = Pt(26)
            r.font.color.rgb = self._rgb(ACCENT_COLOR)

        if subtitle:
            p = self.doc.add_paragraph()
            r = p.add_run(subtitle)
            r.italic         = True
            r.font.size      = Pt(12)
            r.font.color.rgb = self._rgb(SUBTITLE_COLOR)

        self.doc.add_paragraph()

    # ── directive dispatcher ──────────────────────────────────────────────────

    def render_directive(self, name: str, text: str) -> None:
        if name == "figure":
            self._render_figure(text)
            return
        if name == "diagram-placeholder":
            self._render_diagram_placeholder(text)
            return

        cfg = DIRECTIVE_CONFIG.get(name)
        if cfg:
            label, header_fill, content_fill = cfg
            self._generic_box(label, header_fill, content_fill, text)
        else:
            tokens = self._md(text)
            if tokens:
                self.render_tokens(tokens)

    # ── generic box ───────────────────────────────────────────────────────────

    def _generic_box(self, label: str, header_fill: str, content_fill: str, text: str) -> None:
        self._box_header(label, header_fill, "FFFFFF")
        for tok in (self._md(text) or []):
            self._block_in_box(tok, fill=content_fill, border=header_fill)
        self.doc.add_paragraph()

    def _box_header(self, label: str, fill: str, text_color: str) -> None:
        p = self.doc.add_paragraph()
        self._shade_p(p, fill)
        r = p.add_run(f"  {label}")
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = self._rgb(text_color)
        r.font.all_caps  = True

    def _block_in_box(self, tok: dict, fill: str, border: str) -> None:
        t = tok["type"]

        if t in ("paragraph", "block_text"):
            p = self.doc.add_paragraph()
            self._shade_p(p, fill)
            self._left_border(p, border)
            self._indent_p(p, 280)
            self._inline(p, tok.get("children", []))

        elif t == "list":
            ordered = tok.get("ordered", False)
            style   = "List Number" if ordered else "List Bullet"
            for item in tok.get("children", []):
                if item["type"] != "list_item":
                    continue
                inline: list = []
                for child in item.get("children", []):
                    if child["type"] in ("paragraph", "block_text"):
                        inline.extend(child.get("children", []))
                    else:
                        inline.append(child)
                p = self.doc.add_paragraph(style=style)
                self._shade_p(p, fill)
                self._left_border(p, border)
                self._inline(p, inline)

        elif t not in ("newline", "blank_line"):
            self.render_block(tok)

    # ── special element renderers ─────────────────────────────────────────────

    def _render_figure(self, text: str) -> None:
        lines = [ln for ln in text.strip().splitlines() if ln.strip()]
        img_match = None
        caption_parts: list[str] = []

        for line in lines:
            m = IMG_INLINE_RE.search(line.strip())
            if m and img_match is None:
                img_match = m
            else:
                caption_parts.append(line.strip())

        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_match:
            alt, src = img_match.group(1), img_match.group(2)
            ph = PLACEHOLDER_RE.match(src)
            if ph:
                self._image_placeholder(p_img, ph.group(1), alt)
            else:
                r = p_img.add_run(f"[ IMAGE: {alt or src} ]")
                r.font.italic    = True
                r.font.color.rgb = self._rgb("888888")

        if caption_parts:
            p_cap = self.doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_cap.add_run(" ".join(caption_parts))
            r.italic         = True
            r.font.size      = Pt(9)
            r.font.color.rgb = self._rgb("444444")

        self.doc.add_paragraph()

    def _render_diagram_placeholder(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_p(p, "EEEEEE")
        r = p.add_run(f"\n  [ DIAGRAM: {text.strip()} ]\n")
        r.font.name      = "Arial"
        r.font.size      = Pt(10)
        r.font.color.rgb = self._rgb("666666")
        r.italic         = True
        self.doc.add_paragraph()
